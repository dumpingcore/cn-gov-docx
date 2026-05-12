from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree


def _replace_drawing(
    run, rId, docPr_id, docPr_name, cNvPr_id, cNvPr_name, extent_cx, extent_cy
):
    old = run._element.xpath(
        ".//w:drawing",
    )[0]

    new_xml = f"""
<w:drawing
    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
    xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing">
    <wp:anchor distT="0" distB="0" distL="114300" distR="114300" simplePos="0"
        relativeHeight="251658240" behindDoc="0" locked="0" layoutInCell="1"
        allowOverlap="1">
        <wp:simplePos x="0" y="0" />
        <wp:positionH relativeFrom="page">
            <wp:align>center</wp:align>
        </wp:positionH>
        <wp:positionV relativeFrom="paragraph">
            <wp:posOffset>364445</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{extent_cx}" cy="{extent_cy}" />
        <wp:effectExtent l="0" t="0" r="4445" b="2540" />
        <wp:wrapTopAndBottom />
        <wp:docPr id="{docPr_id}" name="{docPr_name}" />
        <wp:cNvGraphicFramePr>
            <a:graphicFrameLocks
                xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                noChangeAspect="1" />
        </wp:cNvGraphicFramePr>
        <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
            <a:graphicData
                uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:pic
                    xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
                    <pic:nvPicPr>
                        <pic:cNvPr id="{cNvPr_id}" name="{cNvPr_name}" />
                        <pic:cNvPicPr />
                    </pic:nvPicPr>
                    <pic:blipFill>
                        <a:blip r:embed="{rId}" />
                        <a:stretch>
                            <a:fillRect />
                        </a:stretch>
                    </pic:blipFill>
                    <pic:spPr>
                        <a:xfrm>
                            <a:off x="0" y="0" />
                            <a:ext cx="{extent_cx}" cy="{extent_cy}" />
                        </a:xfrm>
                        <a:prstGeom prst="rect">
                            <a:avLst />
                        </a:prstGeom>
                    </pic:spPr>
                </pic:pic>
            </a:graphicData>
        </a:graphic>
        <wp14:sizeRelH relativeFrom="margin">
            <wp14:pctWidth>0</wp14:pctWidth>
        </wp14:sizeRelH>
        <wp14:sizeRelV relativeFrom="margin">
            <wp14:pctHeight>0</wp14:pctHeight>
        </wp14:sizeRelV>
    </wp:anchor>
</w:drawing>
    """

    new_element = etree.fromstring(new_xml.encode("utf-8"))

    parent = old.getparent()
    parent.replace(old, new_element)


class CnGovDocx:
    def __init__(self):
        self.doc = Document()

        self._set_a4_layout()
        self._setup_styles()

    def _set_a4_layout(self):
        section = self.doc.sections[0]

        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

    def _setup_styles(self):
        styles = self.doc.styles

        # Title
        title = styles["Title"]
        title.font.name = "方正小标宋简体"
        title.font.size = Pt(22)
        title.font.color.rgb = None  # Text 1
        title.paragraph_format.first_line_indent = Pt(0)
        pPr = title.element.find(qn("w:pPr"))
        if pPr is not None:
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is not None:
                pPr.remove(pBdr)

        # Heading 1
        h1 = styles["Heading 1"]
        h1.font.name = "黑体"
        h1.font.size = Pt(16)
        h1.font.color.rgb = None  # Text 1
        h1.font.bold = False

        h1.paragraph_format.line_spacing = Pt(28)
        h1.paragraph_format.space_before = Pt(0)
        h1.paragraph_format.space_after = Pt(0)
        h1.paragraph_format.keep_with_next = True
        h1.paragraph_format.keep_together = True
        h1.paragraph_format.first_line_indent = Pt(0)

        # Heading 2
        h2 = styles["Heading 2"]
        h2.font.name = "黑体"
        h2.font.size = Pt(16)
        h2.font.color.rgb = None  # Text 1
        h2.font.bold = False

        h2.paragraph_format.line_spacing = Pt(28)
        h2.paragraph_format.space_before = Pt(0)
        h2.paragraph_format.space_after = Pt(0)
        h2.paragraph_format.keep_with_next = True
        h2.paragraph_format.keep_together = True
        h2.paragraph_format.first_line_indent = Pt(0)

        # Heading 2
        h3 = styles["Heading 3"]
        h3.font.name = "黑体"
        h3.font.size = Pt(16)
        h3.font.color.rgb = None  # Text 1
        h3.font.bold = False

        h3.paragraph_format.line_spacing = Pt(28)
        h3.paragraph_format.space_before = Pt(0)
        h3.paragraph_format.space_after = Pt(0)
        h3.paragraph_format.keep_with_next = True
        h3.paragraph_format.keep_together = True
        h3.paragraph_format.first_line_indent = Pt(0)

        # Normal
        normal = styles["Normal"]
        normal.font.name = "仿宋_GB2312"
        normal.font.size = Pt(16)
        normal.font.color.rgb = None

        normal.paragraph_format.line_spacing = Pt(28)
        normal.paragraph_format.first_line_indent = Pt(32)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _apply_font(self, run, font_name: str):
        run.font.name = font_name
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def add_title(self, text: str):
        p = self.doc.add_paragraph(style="Title")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        self._apply_font(run, "方正小标宋简体")

    def add_h1(self, text: str):
        p = self.doc.add_paragraph(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        self._apply_font(run, "黑体")

    def add_h2(self, text: str):
        p = self.doc.add_paragraph(style="Heading 2")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        self._apply_font(run, "黑体")

    def add_h3(self, text: str):
        p = self.doc.add_paragraph(style="Heading 3")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        self._apply_font(run, "黑体")

    def add_text(
        self, text: str, indent_first_line: bool = True, remove_newlines: bool = True
    ):
        if remove_newlines:
            text = "".join([x.strip() for x in text.splitlines()])
        p = self.doc.add_paragraph(style="Normal")

        if not indent_first_line:
            p.paragraph_format.first_line_indent = Pt(0)

        run = p.add_run(text)
        self._apply_font(run, "仿宋_GB2312")

    def add_image_block(self, image_path: str):
        p = self.doc.paragraphs[-1]
        r = p.add_run()
        r.add_picture(image_path)
        blip = r._element.xpath(".//a:blip")[0]
        rId = blip.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        )

        cNvPr = r._element.xpath(".//pic:cNvPr")[0]
        cNvPr_id = cNvPr.get("id")
        cNvPr_name = cNvPr.get("name")

        docPr = r._element.xpath(".//wp:docPr")[0]
        docPr_id = docPr.get("id")
        docPr_name = docPr.get("name")

        extent = r._element.xpath(".//wp:extent")[0]
        extent_cx = int(extent.get("cx"))
        extent_cy = int(extent.get("cy"))

        section = self.doc.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin
        # resize image to fit usable_width
        if extent_cx > usable_width:
            ratio = usable_width / extent_cx
            extent_cy = int(extent_cy * ratio)
            extent_cx = usable_width

        _replace_drawing(
            r, rId, docPr_id, docPr_name, cNvPr_id, cNvPr_name, extent_cx, extent_cy
        )

    def save(self, filepath: str):
        self.doc.save(filepath)
